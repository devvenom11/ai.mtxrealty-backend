from fastapi import FastAPI, HTTPException, Query, UploadFile, File
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.memory import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_community.chat_models import ChatOpenAI
from langchain.memory import ChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from pathlib import Path
from scipy.spatial.distance import cosine
from pydantic import BaseModel
import os
from dotenv import load_dotenv
import shutil
import docx
import whisper

# Load environment variables
load_dotenv()

from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(
    title="Document and Video QA API",
    description="This API allows you to load PDFs, transcriptions, and Google Docs and query them using LangChain and FAISS",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins, but you can specify specific domains
    allow_credentials=True,
    allow_methods=["*"],  # Allow all methods (GET, POST, etc.)
    allow_headers=["*"],  # Allow all headers
)



# Example dictionary with topics and YouTube URLs
TOPICS_YOUTUBE_URLS = {
    "How do I request a CDA?": "https://youtu.be/MmNycgo0KWo",
    "Email signed CDA to title": "https://youtu.be/wqdTUykj_Xc",
    "Here is a good overview of Brokermint.": "https://youtu.be/pErHroeCExY", 
    "How do I create a buyer offer?": "https://youtu.be/zwLWGb34w5M", 
    "Full explanation for a Residential 1-4 Contract": "https://youtu.be/2N6AYFhbQfs", 
    "How do I create a Buyer's Representation Packet": "https://youtu.be/z-M3x0FMQA8",
    "How do I create a Listing Packet?": "https://youtu.be/LoSSI-dc7bs",
    "How do I create a Lease Packet?": "https://youtu.be/SyWfHG3ygOQ", 
    "How do I create an Apartment Packet?": "https://youtu.be/Dz_a_YgKHc4"
}

# Hard-coded paths
SERVER_DATA_FOLDER = "data_pdf"
SERVER_VIDEO_FOLDER = "videos"
SERVER_DRIVE_DOCUMENTS_FOLDER = "data_docs"
SERVER_VIDEO_TRANSCRIPTION_FOLDER = "video_transcripts"

# Initialize variables
store = None
session_store = {}

# Load API key and model name from environment variables
model_name = os.getenv("model_name")
api_key = os.getenv("api_key")

# Initialize OpenAI embeddings
embeddings = OpenAIEmbeddings(model='text-embedding-ada-002', openai_api_key=api_key)

question_answer_pairs = {}

# Define the input model for the question-answer endpoint
class QuestionAnswerInput(BaseModel):
    question: str
    answer: str

# Endpoint to add question-answer pairs to the dictionary
@app.post("/add_question_answers/")
def add_question_answers(input_data: QuestionAnswerInput):
    question = input_data.question.strip()
    answer = input_data.answer.strip()

    # Check if the question already exists in the dictionary
    if question in question_answer_pairs:
        raise HTTPException(status_code=400, detail="Question already exists.")

    # Add the question-answer pair to the dictionary
    question_answer_pairs[question] = answer
    return {"message": "Question-answer pair added successfully.", "question": question, "answer": answer}
# Initialize FAISS vector store
def initialize_vector_store(documents):
    global store
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    texts = text_splitter.split_documents(documents)
    store = FAISS.from_documents(texts, embeddings)
    store.save_local("./vectorstore")

def get_session_history(session_id: str) -> ChatMessageHistory:
    if session_id not in session_store:
        session_store[session_id] = ChatMessageHistory()
    return session_store[session_id]

@app.get("/")
def read_root():
    return {"message": "Welcome to the Document and Video QA API"}

# API to load data from PDF, transcriptions, and documents
@app.post("/load_data/")
def load_data():
    documents = []
    pdf_files_processed = 0
    docx_count = 0
    transcription_count = 0

    # Load PDFs from hard-coded path
    pdf_files = list(Path(SERVER_DATA_FOLDER).glob("*.pdf"))
    if pdf_files:
        pdf_files_processed = len(pdf_files)
        for pdf_file in pdf_files:
            loader = PyPDFLoader(str(pdf_file))
            loaded_docs = loader.load()
            documents.extend(loaded_docs)

    # Load transcriptions from hard-coded path
    transcription_files = list(Path(SERVER_VIDEO_TRANSCRIPTION_FOLDER).glob("*.txt"))
    if transcription_files:
        for transcription_file in transcription_files:
            with open(transcription_file, 'r') as file:
                content = file.read()
            documents.append(Document(page_content=content, metadata={"source": str(transcription_file)}))
            transcription_count += 1

    # Load Google Docs (.docx) from hard-coded path
    docx_files = list(Path(SERVER_DRIVE_DOCUMENTS_FOLDER).glob("*.docx"))
    if docx_files:
        for docx_file in docx_files:
            doc = docx.Document(docx_file)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            documents.append(Document(page_content=content, metadata={"source": str(docx_file)}))
            docx_count += 1

    if not documents:
        raise HTTPException(status_code=404, detail="No documents found in the specified folders.")

    initialize_vector_store(documents)

    return {
        "message": "Processed documents and stored in FAISS.",
        "pdf_count": pdf_files_processed,
        "docx_count": docx_count,
        "transcription_count": transcription_count
    }

# API to add a new YouTube URL with a topic
@app.post("/add_youtube_url/")
def add_youtube_url(topic: str, url: str):
    if topic in TOPICS_YOUTUBE_URLS:
        raise HTTPException(status_code=400, detail="Topic already exists.")
    
    TOPICS_YOUTUBE_URLS[topic] = url
    return {"message": f"Added YouTube URL for topic: {topic}"}

# API to add documents to the PDF folder
@app.post("/add_document/")
async def add_document(file: UploadFile = File(...)):
    Path(SERVER_DATA_FOLDER).mkdir(parents=True, exist_ok=True)
    file_path = Path(SERVER_DATA_FOLDER) / file.filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    return {"message": f"Document {file.filename} added to {SERVER_DATA_FOLDER}."}

# API to add videos to the video folder
@app.post("/add_video/")
async def add_video(file: UploadFile = File(...)):
    Path(SERVER_VIDEO_FOLDER).mkdir(parents=True, exist_ok=True)
    file_path = Path(SERVER_VIDEO_FOLDER) / file.filename
    with open(file_path, "wb") as f:
        shutil.copyfileobj(file.file, f)
    return {"message": f"Video {file.filename} added to {SERVER_VIDEO_FOLDER}."}

@app.post("/transcribe_all_videos/")
def transcribe_all_videos():
    # Ensure the directories exist
    Path(SERVER_VIDEO_FOLDER).mkdir(parents=True, exist_ok=True)
    Path(SERVER_VIDEO_TRANSCRIPTION_FOLDER).mkdir(parents=True, exist_ok=True)

    # Load Whisper model
    whisper_model = whisper.load_model("base")

    # List all video files in the video folder
    video_files = list(Path(SERVER_VIDEO_FOLDER).glob("*"))
    if not video_files:
        raise HTTPException(status_code=404, detail="No video files found in the specified folder.")

    transcriptions = []

    for video_file in video_files:
        # Transcribe the video
        transcription = whisper_model.transcribe(str(video_file))

        # Save the transcription
        transcription_path = Path(SERVER_VIDEO_TRANSCRIPTION_FOLDER) / (video_file.stem + ".txt")
        with open(transcription_path, "w") as f:
            f.write(transcription["text"])

        transcriptions.append({
            "video_file": str(video_file),
            "transcription_file": str(transcription_path)
        })

    return {
        "message": "Transcription completed for all videos.",
        "transcriptions": transcriptions
    }

from scipy.spatial.distance import cosine

@app.get("/query/")
def query_pgvector(query: str, session_id: str = Query(default="0")):
    if store is None:
        raise HTTPException(status_code=500, detail="Vector store is not initialized. Load PDFs first.")

    # Embedding for the query
    query_embedding = embeddings.embed_query(query)

    # Step 1: Check the question-answer dictionary for a matching or similar question
    threshold = 0.8  # 80% similarity threshold
    max_similarity = 0
    best_match = None

    for question, answer in question_answer_pairs.items():
        question_embedding = embeddings.embed_query(question)
        similarity = 1 - cosine(query_embedding, question_embedding)

        if similarity > max_similarity:
            max_similarity = similarity
            best_match = (question, answer)

    # If a similar question is found in the dictionary, return its answer and fetch YouTube URL if needed
    if max_similarity >= threshold and best_match:
        question, answer = best_match
        # Check if there's a relevant YouTube topic
        youtube_url = None
        for topic, topic_url in TOPICS_YOUTUBE_URLS.items():
            topic_embedding = embeddings.embed_query(topic)
            topic_similarity = 1 - cosine(query_embedding, topic_embedding)
            if topic_similarity >= threshold:
                youtube_url = topic_url
                break

        return {
            "query": query,
            "similar_question": question,
            "answer": answer,
            "youtube_url": youtube_url
        }

    # Step 2: If no matching question, proceed with RAG chain for answer generation
    llm = ChatOpenAI(api_key=os.getenv("api_key"), model_name="gpt-4-turbo", temperature=0)
    contextualize_q_system_prompt = (
        "Given a chat history and the latest user question "
        "which might reference context in the chat history, "
        "formulate a standalone question which can be understood "
        "without the chat history. Do NOT answer the question, "
        "just reformulate it if needed and otherwise return it as is."
    )
    contextualize_q_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", contextualize_q_system_prompt),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ]
    )

    # Create the history-aware retriever
    history_aware_retriever = create_history_aware_retriever(
        llm, store.as_retriever(), contextualize_q_prompt
    )

    # Create the QA prompt template
    system_prompt = (
        "You are an assistant for question-answering tasks. "
        "Use the following pieces of retrieved context to answer "
        "the question. If you don't know the answer, say that you "
        "don't know. Use three sentences maximum and keep the "
        "answer concise."
        "\n\n"
        "{context}"
    )
    qa_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", system_prompt),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}"),
        ]
    )

    # Create the chain for question answering
    question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)

    # Combine everything into a retrieval chain
    rag_chain = create_retrieval_chain(history_aware_retriever, question_answer_chain)

    # Use RunnableWithMessageHistory to manage session history
    conversational_rag_chain = RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key="input",
        history_messages_key="chat_history",
        output_messages_key="answer",
    )

    # Handle new session case
    if session_id == "0":
        session_id = os.urandom(8).hex()  # Generate a new session ID

    # Invoke the chain with the query and session ID
    result = conversational_rag_chain.invoke(
        {"input": query},
        config={"configurable": {"session_id": session_id}}
    )

    # Check for YouTube URL based on similarity with topics
    youtube_url = None
    for topic, topic_url in TOPICS_YOUTUBE_URLS.items():
        topic_embedding = embeddings.embed_query(topic)
        topic_similarity = 1 - cosine(query_embedding, topic_embedding)
        if topic_similarity >= threshold:
            youtube_url = topic_url
            break

    return {
        "query": query,
        "session_id": session_id,
        "answer": result["answer"],
        "youtube_url": youtube_url
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=7000)


    # scp -i ~/Downloads/mtx_ai.pem -r /home/prixite/Documents/mtx-realty-chatbot/drive_documents ubuntu@18.218.37.228:/ai.mtxrealty-backend/

